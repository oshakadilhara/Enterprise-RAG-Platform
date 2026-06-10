"""Document text extraction for PDF, DOCX, TXT, CSV."""

import io
from abc import ABC, abstractmethod
from dataclasses import dataclass

import chardet
import pandas as pd
from docx import Document as DocxDocument
from pypdf import PdfReader


@dataclass
class ExtractedPage:
    content: str
    page_number: int | None = None


@dataclass
class ExtractedDocument:
    pages: list[ExtractedPage]
    metadata: dict


class BaseExtractor(ABC):
    @abstractmethod
    async def extract(self, file_content: bytes, file_name: str) -> ExtractedDocument:
        pass


class PDFExtractor(BaseExtractor):
    async def extract(self, file_content: bytes, file_name: str) -> ExtractedDocument:
        reader = PdfReader(io.BytesIO(file_content))
        pages = []
        for i, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(ExtractedPage(content=text.strip(), page_number=i))
        return ExtractedDocument(
            pages=pages,
            metadata={"page_count": len(reader.pages), "file_type": "pdf"},
        )


class DOCXExtractor(BaseExtractor):
    async def extract(self, file_content: bytes, file_name: str) -> ExtractedDocument:
        doc = DocxDocument(io.BytesIO(file_content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)
        return ExtractedDocument(
            pages=[ExtractedPage(content=content, page_number=1)],
            metadata={"paragraph_count": len(paragraphs), "file_type": "docx"},
        )


class TXTExtractor(BaseExtractor):
    async def extract(self, file_content: bytes, file_name: str) -> ExtractedDocument:
        detected = chardet.detect(file_content)
        encoding = detected.get("encoding", "utf-8") or "utf-8"
        content = file_content.decode(encoding, errors="replace")
        return ExtractedDocument(
            pages=[ExtractedPage(content=content.strip(), page_number=1)],
            metadata={"encoding": encoding, "file_type": "txt"},
        )


class CSVExtractor(BaseExtractor):
    async def extract(self, file_content: bytes, file_name: str) -> ExtractedDocument:
        detected = chardet.detect(file_content)
        encoding = detected.get("encoding", "utf-8") or "utf-8"
        df = pd.read_csv(io.BytesIO(file_content), encoding=encoding)
        rows = []
        for idx, row in df.iterrows():
            row_text = " | ".join(f"{col}: {val}" for col, val in row.items() if pd.notna(val))
            rows.append(ExtractedPage(content=row_text, page_number=int(idx) + 1))
        summary = f"CSV with {len(df)} rows and columns: {', '.join(df.columns.tolist())}"
        return ExtractedDocument(
            pages=[ExtractedPage(content=summary, page_number=0)] + rows,
            metadata={"row_count": len(df), "columns": df.columns.tolist(), "file_type": "csv"},
        )


EXTRACTORS: dict[str, BaseExtractor] = {
    "pdf": PDFExtractor(),
    "docx": DOCXExtractor(),
    "txt": TXTExtractor(),
    "csv": CSVExtractor(),
}


def get_extractor(file_type: str) -> BaseExtractor:
    extractor = EXTRACTORS.get(file_type.lower())
    if not extractor:
        raise ValueError(f"Unsupported file type: {file_type}")
    return extractor
